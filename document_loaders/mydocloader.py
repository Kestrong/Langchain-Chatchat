from curses.ascii import isdigit

from langchain.document_loaders.unstructured import UnstructuredFileLoader
from typing import List
import tqdm


class RapidOCRDocLoader(UnstructuredFileLoader):
    def _get_elements(self) -> List:
        def doc2text(filepath):
            from docx.table import _Cell, Table
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            from docx.text.paragraph import Paragraph
            from docx import Document, ImagePart
            from PIL import Image
            from io import BytesIO
            import numpy as np
            from rapidocr_onnxruntime import RapidOCR
            ocr = RapidOCR()
            doc = Document(filepath)
            resp = ""

            def iter_block_items(parent):
                from docx.document import Document
                if isinstance(parent, Document):
                    parent_elm = parent.element.body
                elif isinstance(parent, _Cell):
                    parent_elm = parent._tc
                else:
                    raise ValueError("RapidOCRDocLoader parse fail")

                for child in parent_elm.iterchildren():
                    if isinstance(child, CT_P):
                        yield Paragraph(child, parent)
                    elif isinstance(child, CT_Tbl):
                        yield Table(child, parent)

            b_unit = tqdm.tqdm(total=len(doc.paragraphs)+len(doc.tables),
                               desc="RapidOCRDocLoader block index: 0")
            for i, block in enumerate(iter_block_items(doc)):
                b_unit.set_description(
                    "RapidOCRDocLoader  block index: {}".format(i))
                b_unit.refresh()
                if isinstance(block, Paragraph):
                    resp += block.text.strip() + "\n"
                    images = block._element.xpath('.//pic:pic')  # 获取所有图片
                    for image in images:
                        for img_id in image.xpath('.//a:blip/@r:embed'):  # 获取图片id
                            part = doc.part.related_parts[img_id]  # 根据图片id获取对应的图片
                            if isinstance(part, ImagePart):
                                image = Image.open(BytesIO(part._blob))
                                result, _ = ocr(np.array(image))
                                if result:
                                    ocr_result = [line[1] for line in result]
                                    resp += "\n".join(ocr_result)
                elif isinstance(block, Table):
                    row_cells, column_cells = [], []
                    index = []
                    width, length = len(block.columns), len(block.rows)
                    k = 0
                    # 解决单元格合并后读取出现重复问题
                    for row in block.rows:
                        for cell in row.cells:
                            if cell not in row_cells:
                                index.append([k // width, k % width])
                                row_cells.append(cell)
                            k += 1
                    k = 0
                    for column in block.columns:
                        for cell in column.cells:
                            if cell not in column_cells:
                                column_cells.append(cell)
                            elif [k % length, k // length] in index:
                                index.remove([k % length, k // length])
                            k += 1
                    for rowId in range(length):
                        for colId in range(width):
                            if [rowId, colId] in index:
                                for paragraph in block.rows[rowId].cells[colId].paragraphs:
                                    context = paragraph.text.strip()
                                    if isdigit(resp[-1]) and isdigit(context[0]):
                                        resp += context + " "
                                    else:
                                        resp += context + "\n"
                b_unit.update(1)
            return resp

        text = doc2text(self.file_path)
        from unstructured.partition.text import partition_text
        return partition_text(text=text, **self.unstructured_kwargs)


if __name__ == '__main__':
    loader = RapidOCRDocLoader(file_path="../tests/samples/ocr_test.docx")
    docs = loader.load()
    print(docs)
